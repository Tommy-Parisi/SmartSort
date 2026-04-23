from pathlib import Path
from ...core.utils import token_cap
from ..identity_utils import clean_filename_stem

_HEADING_STYLES = {"heading 1", "heading 2", "heading 3", "title"}
_HEADING_THRESHOLD = 3   # docs with fewer headings are treated as freeform body text
_TARGET_WORDS = 150      # matches token_cap default


class DocxExtractorAgent:
    def extract(self, file_path: str) -> str:
        try:
            from docx import Document
            doc = Document(file_path)
        except Exception:
            return clean_filename_stem(file_path)

        headings, body = [], []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or "").lower() if para.style else ""
            if any(h in style_name for h in _HEADING_STYLES):
                headings.append(text)
            else:
                body.append(text)

        if len(headings) >= _HEADING_THRESHOLD:
            # Structured doc — headings give good cluster signal; add first body para for colour
            parts = headings[:5]
            for bp in body:
                if len(bp) > 30:
                    parts.append(bp)
                    break
        else:
            # Freeform doc (essay, submission, cover letter) — pull body text directly
            parts = headings[:]   # keep any sparse headings as context
            word_count = sum(len(h.split()) for h in parts)
            for bp in body:
                if word_count >= _TARGET_WORDS:
                    break
                parts.append(bp)
                word_count += len(bp.split())

        if not parts:
            # Absolute fallback: first non-empty paragraph
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)
                    break

        return token_cap(" ".join(parts)) if parts else clean_filename_stem(file_path)
